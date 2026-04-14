"""
UNSPED Karbon Ayak İzi — Word Rapor Üreticisi v2
TS EN ISO 14064-1:2019 formatında — python-docx
"""
import sys, json, io, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Veri ──────────────────────────────────────────────────────────
data_path = sys.argv[1] if len(sys.argv) > 1 else "report_data.json"
out_path  = sys.argv[2] if len(sys.argv) > 2 else "UNSPED_Karbon_Raporu.docx"
with open(data_path, encoding="utf-8") as f:
    D = json.load(f)

years        = sorted(D["years"])  # Her zaman 2023→2024→2025 sırası
scope_totals = D["scope_totals"]
cats         = D["cats"]
per_capita   = D.get("per_capita", [])
detail       = D.get("detail", {})
generated_at = D.get("generated_at", "")
selected     = D.get("selected_sections", [])
latest_year  = years[-1]
prev_year    = years[-2] if len(years) >= 2 else None
latest_d     = scope_totals.get(latest_year, {})
prev_d       = scope_totals.get(prev_year, {}) if prev_year else {}
latest_cats  = cats.get(latest_year, {})

# ── Renkler ───────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TXT   = RGBColor(0x59, 0x59, 0x59)
ORANGE_TXT = RGBColor(0x84, 0x3C, 0x0C)
RED_COL    = RGBColor(0xC0, 0x00, 0x00)
GREEN_COL  = RGBColor(0x1D, 0x9E, 0x75)

HDR_BG  = "1F4E79"
HDR2_BG = "2E75B6"
GREY_BG = "F2F2F2"
YLW_BG  = "FFF2CC"
ORG_BG  = "FCE4D6"
LBL_BG  = "D6E4F0"
GRN_BG  = "E2EFDA"
ORG_BG2 = "FFF8EB"

# ── Sayfa ─────────────────────────────────────────────────────────
doc     = Document()
section = doc.sections[0]
section.page_width   = Cm(21.0)
section.page_height  = Cm(29.7)
section.left_margin  = section.right_margin  = Cm(2.5)
section.top_margin   = section.bottom_margin = Cm(2.0)
doc.styles["Normal"].font.name = "Arial"
doc.styles["Normal"].font.size = Pt(10)

# ── Hücre yardımcıları ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper()); tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"),color)
        tcB.append(el)
    tcPr.append(tcB)

def cell_write(cell, text, bold=False, color=None, align="left",
               bg=None, sz=9, italic=False, colspan=1):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg: set_cell_bg(cell, bg)
    set_cell_borders(cell)
    para = cell.paragraphs[0]; para.clear()
    run  = para.add_run(str(text) if text is not None else "—")
    run.font.name = "Arial"; run.font.size = Pt(sz)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    aligns = {"center": WD_ALIGN_PARAGRAPH.CENTER,
              "right":  WD_ALIGN_PARAGRAPH.RIGHT,
              "left":   WD_ALIGN_PARAGRAPH.LEFT}
    para.alignment = aligns.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in [("top","80"),("bottom","80"),("left","120"),("right","120")]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), val); el.set(qn("w:type"), "dxa"); mar.append(el)
    tcPr.append(mar)

def hdr(cell, text, align="center"):
    cell_write(cell, text, bold=True, color=WHITE, align=align, bg=HDR_BG, sz=9)
def hdr2(cell, text, align="center"):
    cell_write(cell, text, bold=True, color=WHITE, align=align, bg=HDR2_BG, sz=9)
def total_row_cells(row, data):
    for cell, (text, align) in zip(row.cells, data):
        cell_write(cell, text, bold=True, color=ORANGE_TXT, align=align, bg=YLW_BG)
def grand_total_cells(row, data):
    for cell, (text, align) in zip(row.cells, data):
        cell_write(cell, text, bold=True, color=WHITE, align=align, bg=HDR_BG)
def alt_bg(i): return GREY_BG if i % 2 == 1 else None

def make_table(doc, rows, cols, col_widths=None):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    return tbl

# ── Paragraf yardımcıları ─────────────────────────────────────────
def add_h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(14)
    r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"2E75B6")
    pBdr.append(bot); pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(12)
    r.font.bold = True; r.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(doc, text, bold=False, color=None, indent=False, sz=10):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(sz)
    r.font.bold = bold
    if color: r.font.color.rgb = color
    if indent: p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph(); r = p.add_run(f"Not: {text}")
    r.font.name = "Arial"; r.font.size = Pt(8)
    r.font.italic = True; r.font.color.rgb = GREY_TXT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)

def add_space(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pt)
    p.paragraph_format.space_after  = Pt(0)

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"• {text}")
    r.font.name = "Arial"; r.font.size = Pt(10)
    if color: r.font.color.rgb = color

# ── Format ────────────────────────────────────────────────────────
def fmtN(v):
    if v is None or v == 0: return "—"
    try:
        f = float(v)
        if f == 0: return "—"
        return f"{f:,.2f}".replace(",","X").replace(".",",").replace("X",".")
    except: return "—"

def fmtA(v):
    """Faaliyet verisi — tam sayı veya ondalık"""
    if v is None or v == 0: return "—"
    try:
        f = float(v)
        if f == int(f): return f"{int(f):,}".replace(",",".")
        return f"{f:,.2f}".replace(",","X").replace(".",",").replace("X",".")
    except: return "—"

def pct_change(a, b):
    if not b or float(b) == 0: return "—"
    d = (float(a) - float(b)) / float(b) * 100
    arrow = "▲" if d > 0 else "▼"
    return f"{arrow} %{abs(d):.0f}"

def pct_color(a, b):
    if not b or float(b) == 0: return GREY_TXT
    return RED_COL if float(a) > float(b) else GREEN_COL

# ── DB'den gelen veriler yardımcısı ──────────────────────────────
def get_stationary(y):
    return detail.get(str(y), {}).get("stationary", [])

def get_mobile(y):
    return detail.get(str(y), {}).get("mobile", [])

def get_refrigerant(y):
    return detail.get(str(y), {}).get("refrigerant", [])

def fuel_co2e(recs, fuels):
    return sum(r["co2e"] for r in recs
               if any(f in (r.get("fuel","") or "").lower() for f in fuels))

# ═══════════════════════════════════════════════════════════════════
# KAPAK
# ═══════════════════════════════════════════════════════════════════
def add_cover(doc):
    for _ in range(8): add_space(doc, 14)
    for text, sz, col, bold in [
        ("UNSPED", 36, DARK_BLUE, True),
        ("Gümrük Müşavirliği ve Lojistik Hizmetler A.Ş.", 13, MID_BLUE, False),
    ]:
        p = doc.add_paragraph(); r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(sz)
        r.font.bold = bold; r.font.color.rgb = col
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

    add_space(doc, 12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"2E75B6")
    pBdr.append(bot); pPr.append(pBdr)
    add_space(doc, 16)

    for text, sz, col, bold in [
        ("Kurumsal Karbon Ayak İzi", 22, DARK_BLUE, True),
        ("Envanter Raporu",          22, MID_BLUE,  True),
        (f"{latest_year} Yılı",      14, GREY_TXT,  False),
        ("TS EN ISO 14064-1:2019", 10, GREY_TXT, False),
        (f"Raporlama Dönemi: {', '.join(years)}", 10, GREY_TXT, False),
        (f"Hazırlanma Tarihi: {generated_at}", 9, GREY_TXT, False),
    ]:
        p = doc.add_paragraph(); r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(sz)
        r.font.bold = bold; r.font.color.rgb = col
        r.font.italic = (not bold and sz <= 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# GİRİŞ
# ═══════════════════════════════════════════════════════════════════
def add_intro(doc):
    add_h1(doc, "GİRİŞ")
    add_body(doc, "UNSPED Gümrük Müşavirliği ve Lojistik Hizmetler A.Ş. kuruluşundan bu yana müşterilerin ihtiyaçlarına beklentilerinin üzerinde yanıt verecek şekilde sürekli olarak teknolojik altyapı ve yeniliklere öncelik vermiştir.")
    add_space(doc)
    add_body(doc, "UNSPED olarak, çevresel sürdürülebilirlik ve sorumluluk bilinciyle hareket etmek amacıyla karbon ayak izinin doğru hesaplanması ve yönetilmesi büyük önem taşımaktadır. Bu Kurumsal Karbon Ayak İzi Envanter Raporu, UNSPED Genel Merkezi sera gazı emisyonlarını detaylı bir şekilde ortaya koymak amacıyla hazırlanmıştır.")
    add_space(doc)
    add_body(doc, "Bu rapor, uygunluk, bütünlük, tutarlılık, doğruluk ve şeffaflık ilkeleri uygulanarak hazırlanmıştır.")
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# BÖLÜM 1: TEKNİK BİLGİLER
# (1.5 metodoloji burada SADECE kısa metin — tablolar Bölüm 8'de)
# ═══════════════════════════════════════════════════════════════════
def add_technical(doc):
    add_h1(doc, "BÖLÜM 1: RAPORUN AMAÇLARI VE TEKNİK BİLGİLER")

    add_h2(doc, "1.1. Şirket Profili")
    add_body(doc, "UNSPED Gümrük Müşavirliği ve Lojistik Hizmetler A.Ş., müşterilerinin dış ticaret ve gümrük işlemlerindeki durumlarını analiz ederek risk yönetimi, maliyet optimizasyonu ve lojistik çözümleri sunmaktadır.")

    add_h2(doc, "1.2. Raporun Amacı")
    add_body(doc, "Bu rapor TS EN ISO 14064-1:2019 standardı şartlarını esas alarak hazırlanmıştır.")

    add_h2(doc, "1.3. Raporlama Periyodu")
    add_body(doc, f"Bu rapor {latest_year} yılı sera gazı emisyonlarını kapsar. Temel yıl: {years[0]}.")

    add_h2(doc, "1.4. Temel Yıl")
    add_body(doc, f"TS EN ISO 14064-1:2019 kapsamında temel yıl {years[0]} olarak belirlenmiştir.")

    add_h2(doc, "1.5. Hesaplama Metodolojisi")
    add_body(doc, "Bu raporda emisyon faktörü bazlı hesaplama metodolojisi kullanılmıştır. Detaylı emisyon faktörü ve NKD tabloları Bölüm 8'de yer almaktadır. Kullanılan başlıca kaynaklar:")
    add_space(doc, 4)
    for madde in [
        "IPCC 2006 Guidelines for National GHG Inventories — yakıt NKD ve EF değerleri",
        "IPCC AR6 (2021) — GWP-100 değerleri",
        "TC ETKB EVÇED — Türkiye ulusal elektrik şebekesi emisyon faktörleri",
        "UK DEFRA GHG Conversion Factors — taşımacılık emisyon faktörleri",
        "US EPA USEEIO v1.2 — satın alma bazlı tedarik zinciri emisyon faktörleri (kg CO₂e/2021 USD)",
    ]:
        add_bullet(doc, madde)

    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# BÖLÜM 2-3: SINIRLAR
# ═══════════════════════════════════════════════════════════════════
def add_boundaries(doc):
    add_h1(doc, "BÖLÜM 2: KURULUŞ SINIRLARI VE RAPORLAMA")
    add_body(doc, "Emisyon konsolidasyonu için Operasyonel Kontrol Yaklaşımı uygulanmıştır.")
    add_space(doc)
    add_caption(doc, "Tablo 1. Envantere Dahil Edilen Tesis")
    tbl = make_table(doc, 2, 2, [8.0, 7.0])
    hdr(tbl.rows[0].cells[0], "Tesis Adı", "left")
    hdr(tbl.rows[0].cells[1], "Adres", "left")
    cell_write(tbl.rows[1].cells[0], "UNSPED Gümrük Müşavirliği ve Lojistik Hizmetler A.Ş. Genel Merkezi")
    cell_write(tbl.rows[1].cells[1], "Mahmutbey Mahallesi 2655 Sokak No:1-3-21, Bağcılar, İstanbul")

    add_h1(doc, "BÖLÜM 3: RAPORLAMA SINIRLARI")
    add_body(doc, "Raporlama sınırları TS EN ISO 14064-1 kapsamında aşağıdaki emisyon kategorilerini içermektedir:")
    add_space(doc)
    for kapsam, aciklama in [
        ("Kapsam 1 (Sınıf 1)", "Sabit yakma (doğalgaz, motorin), mobil yanma (benzin, dizel), soğutucu gaz ve yangın söndürme sızıntıları/şarjları"),
        ("Kapsam 2 (Sınıf 2)", "Satın alınan şebeke elektriği — konum temelli yöntem"),
        ("Kapsam 3 (Sınıf 3)", "Personel işe gidiş-geliş (3.3), iş seyahatleri (3.5)"),
        ("Kapsam 3 (Sınıf 4)", "Satın alınan mallar (4.1), sermaye varlıkları (4.2), atık bertarafı (4.3), hizmet alımları (4.5)"),
        ("Kapsam 3 (Sınıf 6)", "Elektrik iletim ve dağıtım kayıpları (6.1) — konum temelli"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{kapsam}: ")
        r1.font.bold = True; r1.font.color.rgb = DARK_BLUE; r1.font.size = Pt(10)
        r2 = p.add_run(aciklama); r2.font.size = Pt(10)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# BÖLÜM 4: EMİSYONLAR
# ═══════════════════════════════════════════════════════════════════

def _year_cols(n_extra=0):
    """Yıl sütunu genişliği — yıl sayısına göre dinamik."""
    return min(2.5, 8.0 / max(len(years), 1))

def _tbl_year_hdrs(tbl, prefix_count):
    for yi, y in enumerate(years):
        hdr(tbl.rows[0].cells[prefix_count + yi],
            f"{y}\n(tCO₂e)", "center")

def add_scope1(doc):
    add_h1(doc, "BÖLÜM 4: SERA GAZI EMİSYONLARI")
    add_h2(doc, "4.1. Kapsam 1: Doğrudan Sera Gazı Emisyonları")
    add_body(doc, "Kapsam 1 emisyonları şirketin sahip olduğu veya kontrol ettiği kaynaklardan doğrudan salınan sera gazlarını kapsar.")

    # ── 4.1.1 Sabit Yanma ─────────────────────────────────────────
    add_h3(doc, "4.1.1. Kapsam 1.1 — Sabit Yanma")
    yc = _year_cols()
    add_caption(doc, f"Tablo 2. Kapsam 1.1 Sabit Yanma Emisyonları (ton CO₂e)")
    # Sütunlar: Kod | Kaynak | Yakıt | Birim | FV(yıllar) | CO2e(yıllar)
    cw = [1.0, 3.8, 1.8, 1.0] + [yc]*len(years)*2
    tbl = make_table(doc, 6, 4 + len(years)*2, cw)
    # Başlıklar
    for i, txt in enumerate(["Kod","Emisyon Kaynağı","Yakıt","Birim"]):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==1 else "center")
    for yi, y in enumerate(years):
        hdr2(tbl.rows[0].cells[4+yi],         f"{y} FV",    "center")
        hdr2(tbl.rows[0].cells[4+len(years)+yi], f"{y} tCO₂e","center")

    fuel_map_11 = [
        ("1.1.1","Kombi","Doğalgaz","m³", ["natural_gas","dogalgaz","gaz"]),
        ("1.1.2","Jeneratörler","Motorin","L", ["diesel","motorin","mazot"]),
        ("1.1.3","Yangın Tatbikatı","Motorin","L", ["diesel","motorin"]),
        ("1.1.4","Yangın Tatbikatı","Benzin","L", ["petrol","benzin","gasoline"]),
    ]
    for ri, (kod, lbl, ykt, birim, fuels) in enumerate(fuel_map_11):
        bg = alt_bg(ri); row = tbl.rows[ri+1]
        cell_write(row.cells[0], kod,   align="center", bg=bg)
        cell_write(row.cells[1], lbl,   bg=bg)
        cell_write(row.cells[2], ykt,   bg=bg)
        cell_write(row.cells[3], birim, align="center", bg=bg)
        for yi, y in enumerate(years):
            recs = get_stationary(y)
            co2e = fuel_co2e(recs, fuels)
            fv   = sum(r.get("value",0) for r in recs
                       if any(f in (r.get("fuel","") or "").lower() for f in fuels))
            cell_write(row.cells[4+yi],          fmtA(fv)   if fv   else "—", align="right", bg=bg)
            cell_write(row.cells[4+len(years)+yi],fmtN(co2e) if co2e else "—", align="right", bg=bg)

    # Toplam satırı
    total_data = [("","left"),("1.1 TOPLAM","left"),("","center"),("","center")]
    for _ in years: total_data.append(("","right"))  # FV toplamı yok
    for y in years: total_data.append((fmtN(cats.get(y,{}).get("1.1",0)),"right"))
    total_row_cells(tbl.rows[5], total_data)
    add_space(doc)

    # ── 4.1.2 Mobil Yanma ─────────────────────────────────────────
    add_h3(doc, "4.1.2. Kapsam 1.2 — Mobil Yanma (Araç Yakıtları)")
    add_caption(doc, f"Tablo 3. Kapsam 1.2 Mobil Yanma Emisyonları (ton CO₂e)")
    cw2 = [1.0, 3.8, 1.8, 1.0] + [yc]*len(years)*2
    tbl2 = make_table(doc, 4, 4+len(years)*2, cw2)
    for i, txt in enumerate(["Kod","Emisyon Kaynağı","Yakıt","Birim"]):
        hdr(tbl2.rows[0].cells[i], txt, "left" if i==1 else "center")
    for yi, y in enumerate(years):
        hdr2(tbl2.rows[0].cells[4+yi],          f"{y} FV (L)","center")
        hdr2(tbl2.rows[0].cells[4+len(years)+yi],f"{y} tCO₂e","center")

    mobile_rows = [
        ("1.2.1","Binek Araçlar","Benzin","L",["petrol","benzin","gasoline"]),
        ("1.2.2","Binek Araçlar","Dizel", "L",["diesel","motorin","mazot"]),
    ]
    for ri, (kod,lbl,ykt,birim,fuels) in enumerate(mobile_rows):
        bg = alt_bg(ri); row = tbl2.rows[ri+1]
        cell_write(row.cells[0], kod,   align="center", bg=bg)
        cell_write(row.cells[1], lbl,   bg=bg)
        cell_write(row.cells[2], ykt,   bg=bg)
        cell_write(row.cells[3], birim, align="center", bg=bg)
        for yi, y in enumerate(years):
            recs = get_mobile(y)
            co2e = fuel_co2e(recs, fuels)
            fv   = sum(r.get("value",0) for r in recs
                       if any(f in (r.get("fuel","") or "").lower() for f in fuels))
            cell_write(row.cells[4+yi],          fmtA(fv)   if fv   else "—", align="right", bg=bg)
            cell_write(row.cells[4+len(years)+yi],fmtN(co2e) if co2e else "—", align="right", bg=bg)

    t12 = [("","left"),("1.2 TOPLAM","left"),("","center"),("","center")]
    for _ in years: t12.append(("","right"))
    for y in years: t12.append((fmtN(cats.get(y,{}).get("1.2",0)),"right"))
    total_row_cells(tbl2.rows[3], t12)
    add_space(doc)

    # ── 4.1.3 Soğutucu Gazlar — GERÇEK VERİ ───────────────────────
    add_h3(doc, "4.1.3. Kapsam 1.4 — Soğutucu Gazlar ve Yangın Söndürme Sistemleri")
    add_body(doc, "GWP değerleri IPCC AR6 2021 raporundan alınmıştır. Tüm sızıntı ve şarj olayları aşağıda raporlanmıştır.")
    add_space(doc)
    add_caption(doc, f"Tablo 4. Kapsam 1.4 Soğutucu Gaz Emisyonları (ton CO₂e)")

    # Tüm yıllardaki unique gazları topla
    all_refs = {}  # gas_type → {y: {co2e, fv, unit}}
    for y in years:
        for r in get_refrigerant(y):
            g = r.get("source") or r.get("fuel") or "Bilinmeyen"
            if g not in all_refs: all_refs[g] = {}
            all_refs[g][y] = {
                "co2e": r.get("co2e", 0),
                "fv":   r.get("value", 0),
                "unit": r.get("unit", "ton")
            }

    if all_refs:
        cw3 = [2.8, 1.5, 1.5] + [yc]*len(years)*2
        n_rows = len(all_refs) + 2  # başlık + veri + toplam
        tbl3 = make_table(doc, n_rows, 3+len(years)*2, cw3)
        for i, txt in enumerate(["Emisyon Kaynağı / Gaz","EF (Oran)","Birim"]):
            hdr(tbl3.rows[0].cells[i], txt, "left" if i==0 else "center")
        for yi, y in enumerate(years):
            hdr2(tbl3.rows[0].cells[3+yi],          f"{y} FV",    "center")
            hdr2(tbl3.rows[0].cells[3+len(years)+yi],f"{y} tCO₂e","center")

        for ri, (gas, year_data) in enumerate(sorted(all_refs.items())):
            bg = alt_bg(ri); row = tbl3.rows[ri+1]
            cell_write(row.cells[0], gas, bg=bg, bold=True)
            cell_write(row.cells[1], "Bkz. Tablo EF", align="center", bg=bg, sz=8)
            # Birim — ilk yıldan al
            unit = next((d["unit"] for d in year_data.values() if d.get("unit")), "ton")
            cell_write(row.cells[2], unit, align="center", bg=bg)
            for yi, y in enumerate(years):
                d = year_data.get(y, {})
                cell_write(row.cells[3+yi],          fmtA(d.get("fv",0))  if d else "—", align="right", bg=bg)
                cell_write(row.cells[3+len(years)+yi],fmtN(d.get("co2e",0)) if d else "—", align="right", bg=bg)

        t14 = [("1.4 TOPLAM","left"),("","center"),("","center")]
        for _ in years: t14.append(("","right"))
        for y in years: t14.append((fmtN(cats.get(y,{}).get("1.4",0)),"right"))
        total_row_cells(tbl3.rows[-1], t14)
    else:
        # Veri yoksa toplam göster
        add_body(doc, f"Kapsam 1.4 toplam: {fmtN(latest_cats.get('1.4',0))} ton CO₂e")

    doc.add_page_break()

def add_scope2(doc):
    add_h2(doc, "4.2. Kapsam 2: Enerji Dolaylı Sera Gazı Emisyonları")
    add_body(doc, "TC Enerji ve Tabii Kaynaklar Bakanlığı EVÇED ulusal şebeke emisyon faktörü kullanılmıştır (konum temelli yöntem).")
    add_space(doc)
    add_caption(doc, "Tablo 5. Kapsam 2.1 Elektrik Tüketimi Emisyonları")
    yc = _year_cols()
    cw = [1.5, 5.5, 1.2] + [yc]*len(years)
    tbl = make_table(doc, 3, 3+len(years), cw)
    for i, txt in enumerate(["Kategori","Emisyon Kaynağı","Birim"]):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==1 else "center")
    _tbl_year_hdrs(tbl, 3)

    row = tbl.rows[1]
    cell_write(row.cells[0], "2.1", align="center")
    cell_write(row.cells[1], "Elektrik Tüketimi (Konum Temelli)")
    cell_write(row.cells[2], "MWh → tCO₂e", align="center", sz=8)
    for yi, y in enumerate(years):
        cell_write(row.cells[3+yi], fmtN(cats.get(y,{}).get("2.1",0)), align="right")

    total_row_cells(tbl.rows[2],
        [("","left"),("KAPSAM 2 TOPLAM","left"),("","center")] +
        [(fmtN(scope_totals.get(y,{}).get("k2",0)),"right") for y in years])
    add_space(doc)

def add_scope3(doc):
    add_h2(doc, "4.3. Kapsam 3: Diğer Dolaylı Emisyonlar")
    add_body(doc, "Önemlilik eşiği olarak %95 kuralı uygulanmıştır.")
    add_space(doc)
    add_caption(doc, "Tablo 6. Kapsam 3 Emisyon Kategorileri (ton CO₂e)")
    yc = _year_cols()
    s3cats = [
        ("3.3","Personelin İşe Gidiş-Gelişleri","IPCC/DEFRA"),
        ("3.5","İş Seyahatleri","EPA/DEFRA"),
        ("4.1","Satın Alınan Ürün ve Malzemeler","EPA USEEIO"),
        ("4.2","Sermaye Niteliğindeki Varlıklar","EPA USEEIO"),
        ("4.3","Atık Bertarafı","DEFRA"),
        ("4.5","Hizmet Alımları","EPA USEEIO"),
        ("6.1","Elektrik T&D Kayıpları","TC ETKB EVÇED"),
    ]
    cw = [1.0, 6.0, 2.3] + [yc]*len(years)
    tbl = make_table(doc, len(s3cats)+2, 3+len(years), cw)
    for i, txt in enumerate(["Sınıf","Emisyon Kaynağı","EF Kaynağı"]):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==1 else "center")
    _tbl_year_hdrs(tbl, 3)

    for ri, (kod,lbl,kaynak) in enumerate(s3cats):
        bg = alt_bg(ri); row = tbl.rows[ri+1]
        cell_write(row.cells[0], kod,    align="center", bg=bg)
        cell_write(row.cells[1], lbl,    bg=bg)
        cell_write(row.cells[2], kaynak, align="center", bg=bg, sz=8)
        for yi, y in enumerate(years):
            cell_write(row.cells[3+yi], fmtN(cats.get(y,{}).get(kod,0)),
                       align="right", bg=bg)

    total_row_cells(tbl.rows[-1],
        [("","left"),("KAPSAM 3 TOPLAM","left"),("","center")] +
        [(fmtN(scope_totals.get(y,{}).get("k3",0)),"right") for y in years])
    add_space(doc)

# ═══════════════════════════════════════════════════════════════════
# FALİYET VERİLERİ — 1.1'den 6.1'e %95 eşiğine kadar
# ═══════════════════════════════════════════════════════════════════
def add_activity_data(doc):
    add_h2(doc, "4.3.1. Faaliyet Verileri — Emisyon Kaynak Detayları")
    add_body(doc, f"Aşağıdaki tablolarda {latest_year} yılı faaliyet verileri emisyon büyüklüğüne göre sıralanmış şekilde sunulmaktadır. Kapsam 3 için %95 önemlilik eşiği esas alınmıştır.")
    add_space(doc)

    # Kapsam 1.1 — Sabit yanma kaynakları
    recs_11 = []
    for y in sorted(years):  # 2023→2024→2025
        for r in get_stationary(y):
            recs_11.append((y, r))

    if recs_11:
        add_h3(doc, "Kapsam 1.1 — Sabit Yanma Faaliyet Verileri")
        cw_a = [1.0, 1.2, 3.5, 1.5, 1.5, 2.0]
        tbl_a = make_table(doc, len(recs_11)+1, 6, cw_a)
        for i, txt in enumerate(["Yıl","Yakıt","Kaynak","FV","Birim","tCO₂e"]):
            hdr(tbl_a.rows[0].cells[i], txt, "left" if i==2 else "center")
        for ri, (y, r) in enumerate(sorted(recs_11, key=lambda x: x[1].get("co2e",0), reverse=True)):
            bg = alt_bg(ri); row = tbl_a.rows[ri+1]
            cell_write(row.cells[0], str(y),                  align="center", bg=bg)
            cell_write(row.cells[1], r.get("fuel",""),         align="center", bg=bg, sz=8)
            cell_write(row.cells[2], r.get("source",""),       bg=bg)
            cell_write(row.cells[3], fmtA(r.get("value",0)),   align="right",  bg=bg)
            cell_write(row.cells[4], r.get("unit",""),         align="center", bg=bg)
            cell_write(row.cells[5], fmtN(r.get("co2e",0)),   align="right",  bg=bg)
        add_space(doc)

    # Kapsam 1.2 — Mobil yanma
    recs_12 = []
    for y in sorted(years):  # 2023→2024→2025
        for r in get_mobile(y):
            recs_12.append((y, r))

    if recs_12:
        add_h3(doc, "Kapsam 1.2 — Mobil Yanma Faaliyet Verileri")
        tbl_b = make_table(doc, len(recs_12)+1, 6, cw_a)
        for i, txt in enumerate(["Yıl","Yakıt","Kaynak","FV","Birim","tCO₂e"]):
            hdr(tbl_b.rows[0].cells[i], txt, "left" if i==2 else "center")
        for ri, (y, r) in enumerate(sorted(recs_12, key=lambda x: x[1].get("co2e",0), reverse=True)):
            bg = alt_bg(ri); row = tbl_b.rows[ri+1]
            cell_write(row.cells[0], str(y),                  align="center", bg=bg)
            cell_write(row.cells[1], r.get("fuel",""),         align="center", bg=bg, sz=8)
            cell_write(row.cells[2], r.get("source",""),       bg=bg)
            cell_write(row.cells[3], fmtA(r.get("value",0)),   align="right",  bg=bg)
            cell_write(row.cells[4], r.get("unit",""),         align="center", bg=bg)
            cell_write(row.cells[5], fmtN(r.get("co2e",0)),   align="right",  bg=bg)
        add_space(doc)

    # Kapsam 3 — %95 eşiğine giren kategoriler
    s3_all = {
        "3.3":"Personel Ulaşım","3.5":"Uçak & Konaklama",
        "4.1":"Satın Alınan Mallar","4.2":"Sermaye Varlıkları",
        "4.3":"Atık","4.5":"Hizmet Alımları","6.1":"T&D Kayıpları"
    }
    s3_total = latest_d.get("k3", 0)
    s3_sorted = sorted(
        [(k, latest_cats.get(k,0)) for k in s3_all],
        key=lambda x: x[1], reverse=True
    )
    cum = 0; critical_cats = []
    for k, v in s3_sorted:
        if s3_total > 0:
            cum += v / s3_total * 100
        critical_cats.append(k)
        if cum >= 95: break

    add_h3(doc, f"Kapsam 3 — %95 Eşiğine Giren Kategoriler ({latest_year})")
    cw_s3 = [1.0, 4.5, 2.0, 2.0, 2.0, 2.0]
    n_rows = len(critical_cats) + 2
    tbl_s3 = make_table(doc, n_rows, 6, cw_s3)
    for i, txt in enumerate(["Sınıf","Emisyon Kaynağı","tCO₂e","Pay (%)","Kümülatif (%)","Durum"]):
        hdr(tbl_s3.rows[0].cells[i], txt, "left" if i==1 else "center")

    cum2 = 0
    for ri, k in enumerate(critical_cats):
        val  = latest_cats.get(k, 0)
        p    = val / s3_total * 100 if s3_total > 0 else 0
        cum2 += p
        bg   = ORG_BG
        row  = tbl_s3.rows[ri+1]
        cell_write(row.cells[0], k,              align="center", bg=bg)
        cell_write(row.cells[1], s3_all.get(k,k),bg=bg)
        cell_write(row.cells[2], fmtN(val),      align="right",  bg=bg)
        cell_write(row.cells[3], f"%{p:.1f}",    align="right",  bg=bg)
        cell_write(row.cells[4], f"%{cum2:.1f}", align="right",  bg=bg)
        cell_write(row.cells[5], "Kritik",       align="center", bg=bg,
                   bold=True, color=ORANGE_TXT)

    total_row_cells(tbl_s3.rows[-1],
        [("","left"),("KAPSAM 3 TOPLAM","left")] +
        [(fmtN(latest_d.get("k3",0)),"right"),
         ("%100","right"),("","right"),("","right")])
    add_space(doc)

# ═══════════════════════════════════════════════════════════════════
# ÖZET TABLO
# ═══════════════════════════════════════════════════════════════════
def add_summary(doc):
    add_h2(doc, "4.4. Kapsamlara Göre Sera Gazı Emisyonları — Özet")
    t = latest_d.get("total", 0); p = prev_d.get("total", 0)
    add_body(doc, f"{latest_year} yılı toplam sera gazı emisyonu {fmtN(t)} tCO₂e olarak hesaplanmıştır." +
             (f" {prev_year}'e göre {pct_change(t,p)} değişim göstermiştir." if prev_year else ""))
    add_space(doc)
    add_caption(doc, "Tablo 7. Yıllara Göre GHG Emisyon Özeti (ton CO₂e)")
    yc = _year_cols()
    chg_col = [2.0] if prev_year else []
    cw = [2.5, 4.5] + [yc]*len(years) + chg_col
    tbl = make_table(doc, 5, 2+len(years)+(1 if prev_year else 0), cw)
    hdrs = ["Kapsam","Açıklama"] + years + ([f"Değişim ({prev_year}→{latest_year})"] if prev_year else [])
    for i, txt in enumerate(hdrs):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==1 else "center")
    for ri, (key, lbl, bg) in enumerate([
        ("k1","Kapsam 1 — Doğrudan Emisyonlar",      LBL_BG),
        ("k2","Kapsam 2 — Enerji Dolaylı Emisyonlar", GRN_BG),
        ("k3","Kapsam 3 — Diğer Dolaylı Emisyonlar",  ORG_BG2),
    ]):
        row = tbl.rows[ri+1]
        cell_write(row.cells[0], f"Kapsam {key[1]}", bold=True, bg=bg, color=DARK_BLUE)
        cell_write(row.cells[1], lbl, bg=bg)
        for yi, y in enumerate(years):
            cell_write(row.cells[2+yi], fmtN(scope_totals.get(y,{}).get(key,0)),
                       align="right", bg=bg)
        if prev_year:
            chg = pct_change(latest_d.get(key,0), prev_d.get(key,0))
            col = pct_color(latest_d.get(key,0), prev_d.get(key,0))
            cell_write(row.cells[2+len(years)], chg, align="center", bg=bg, color=col)

    gc = [("GENEL TOPLAM","left"),("Kapsam 1+2+3","left")]
    for y in years: gc.append((fmtN(scope_totals.get(y,{}).get("total",0)),"right"))
    if prev_year:
        chg = pct_change(latest_d.get("total",0),prev_d.get("total",0))
        gc.append((chg,"center"))
    grand_total_cells(tbl.rows[4], gc)
    add_space(doc)

# ═══════════════════════════════════════════════════════════════════
# GRAFİKLER — chart_style.py ile
# ═══════════════════════════════════════════════════════════════════
def add_charts(doc):
    """Grafikleri chart_style.py modülünden alır — Word için statik PNG."""
    try:
        import matplotlib; matplotlib.use("Agg")
        import sys, os
        # chart_style.py — word_export klasöründen bir üst dizinde
        script_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(script_dir)
        for p in [parent_dir, script_dir]:
            if p not in sys.path: sys.path.insert(0, p)
        from chart_style import make_chart1, make_chart2, make_kpi, fig_to_buf
    except ImportError as e:
        add_body(doc, f"Not: Grafik için matplotlib gerekli. ({e})", color=GREY_TXT)
        return

    add_h2(doc, "4.5. Sera Gazı Emisyonları — Grafiksel Gösterim")

    # Veri hazırla
    ys   = years
    k1v  = [scope_totals.get(y,{}).get("k1",0)    for y in ys]
    k2v  = [scope_totals.get(y,{}).get("k2",0)    for y in ys]
    k3v  = [scope_totals.get(y,{}).get("k3",0)    for y in ys]
    totv = [scope_totals.get(y,{}).get("total",0) for y in ys]

    # KPI kartları
    fig_kpi = make_kpi(ys, k1v, k2v, k3v, totv)
    buf_kpi = fig_to_buf(fig_kpi, dpi=140)
    add_caption(doc, f"Şekil 0. KPI Özeti — {latest_year}")
    doc.add_picture(buf_kpi, width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_space(doc)

    # Grouped bar + Donut
    fig1 = make_chart1(ys, k1v, k2v, k3v, totv)
    buf1 = fig_to_buf(fig1, dpi=140)
    add_caption(doc, f"Şekil 1. Yıllık Kapsam Karşılaştırması ve {latest_year} Emisyon Dağılımı")
    doc.add_picture(buf1, width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_space(doc)

    # Stacked trend
    if len(ys) >= 2:
        fig2 = make_chart2(ys, k1v, k2v, k3v, totv)
        buf2 = fig_to_buf(fig2, dpi=140)
        add_caption(doc, "Şekil 2. Yıllık Toplam Emisyon Trendi ve Değişim Oranları")
        doc.add_picture(buf2, width=Inches(5.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_space(doc)

    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# ETKİ ANALİZİ — sadece metin + liste, tablo yok (feedback 7)
# ═══════════════════════════════════════════════════════════════════
def add_impact(doc):
    add_h2(doc, "4.6. Kapsam 3 Emisyon Analizi — %95 Önemlilik Eşiği")
    s3total = latest_d.get("k3", 0)
    s3labels = {
        "3.1":"Hammadde Sevkiyatı","3.2":"Ürün Sevkiyatı",
        "3.3":"Personel Ulaşım","3.4":"İş Seyahati (Kara)",
        "3.5":"Uçak & Konaklama","4.1":"Satın Alınan Mallar",
        "4.2":"Sermaye Varlıkları","4.3":"Atık Bertarafı",
        "4.4":"Kiralanan Ekipman","4.5":"Hizmet Alımları",
        "6.1":"T&D Kayıpları"
    }
    sorted_s3 = sorted(
        [(k,v,latest_cats.get(k,0)) for k,v in s3labels.items() if latest_cats.get(k,0)>0],
        key=lambda x: x[2], reverse=True
    )

    cum = 0; lines = []
    for k, lbl, val in sorted_s3:
        p = val/s3total*100 if s3total > 0 else 0
        cum += p
        lines.append((k, lbl, val, p, cum, cum <= 95))

    add_body(doc,
        f"{latest_year} yılı Kapsam 3 toplam emisyonu {fmtN(s3total)} tCO₂e olarak "
        f"hesaplanmıştır. Emisyonların %95'i aşağıdaki kategorilerden kaynaklanmaktadır:")
    add_space(doc, 4)

    # Kritik olanlar metin olarak listele
    for k, lbl, val, p, c, is_crit in lines:
        if is_crit:
            p_str = doc.add_paragraph()
            p_str.paragraph_format.left_indent = Cm(1)
            p_str.paragraph_format.space_after  = Pt(4)
            r1 = p_str.add_run(f"• {k} {lbl}: ")
            r1.font.bold = True; r1.font.color.rgb = ORANGE_TXT; r1.font.size = Pt(10)
            r2 = p_str.add_run(f"{fmtN(val)} tCO₂e — %{p:.1f} pay (kümülatif: %{c:.1f})")
            r2.font.size = Pt(10)

    # Eşik sonrası
    remaining = [(k,lbl,val,p) for k,lbl,val,p,c,ic in lines if not ic]
    if remaining:
        add_space(doc, 4)
        add_body(doc, f"Önemlilik eşiğinin (%95) altında kalan kategoriler "
                 f"(toplam %{sum(p for _,_,_,p in remaining):.1f}):", color=GREY_TXT)
        for k, lbl, val, p in remaining:
            add_bullet(doc, f"{k} {lbl}: {fmtN(val)} tCO₂e (%{p:.1f})", color=GREY_TXT)

    add_note(doc, "Bu analizde %95 önemlilik eşiği esas alınmıştır (TS EN ISO 14064-1:2019).")
    add_space(doc)

# ═══════════════════════════════════════════════════════════════════
# KİŞİ BAŞINA
# ═══════════════════════════════════════════════════════════════════
def add_percap(doc):
    if not per_capita: return
    add_h2(doc, "4.7. Kişi Bazlı Emisyon Yoğunluğu")
    add_body(doc, "Spesifik metrik: 12 aylık ortalama çalışan sayısına göre kişi başına emisyon (ton CO₂e/kişi).")
    add_space(doc)
    add_caption(doc, "Tablo 8. Kişi Başına Emisyon Yoğunluğu")
    tbl = make_table(doc, len(per_capita)+1, 8,
                     [1.0, 3.5, 1.5, 1.8, 1.8, 1.8, 2.2, 1.8])
    for i, txt in enumerate(["Yıl","Lokasyon","Personel","Kapsam 1","Kapsam 2","Kapsam 3","Toplam (tCO₂e)","tCO₂e/Kişi"]):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==1 else "center")
    for ri, r in enumerate(per_capita):
        is_t = r.get("is_total",False)
        bg   = YLW_BG if is_t else alt_bg(ri)
        row  = tbl.rows[ri+1]
        cell_write(row.cells[0], str(r.get("year","")),    align="center", bg=bg, bold=is_t)
        cell_write(row.cells[1], r.get("location",""),     bg=bg, bold=is_t)
        cell_write(row.cells[2], str(r.get("headcount","")),align="right",  bg=bg, bold=is_t)
        cell_write(row.cells[3], fmtN(r.get("k1",0)),      align="right",  bg=bg)
        cell_write(row.cells[4], fmtN(r.get("k2",0)),      align="right",  bg=bg)
        cell_write(row.cells[5], fmtN(r.get("k3",0)),      align="right",  bg=bg)
        cell_write(row.cells[6], fmtN(r.get("total",0)),   align="right",  bg=bg, bold=is_t)
        cell_write(row.cells[7], fmtN(r.get("per_cap",0)), align="right",  bg=bg, bold=is_t)
    add_space(doc)

# ═══════════════════════════════════════════════════════════════════
# YILLARA GÖRE GHG
# ═══════════════════════════════════════════════════════════════════
def add_yearly(doc):
    if len(years) < 2: return
    add_h1(doc, "BÖLÜM 5: YILLARA GÖRE GHG EMİSYON VERİLERİ")
    add_body(doc, f"{', '.join(years)} yıllarına ait GHG emisyon verileri karşılaştırmalı olarak sunulmaktadır.")
    add_space(doc)
    detail_rows = [
        ("1.1","Sabit Yanma","Kapsam 1"),
        ("1.2","Mobil Yanma (Benzin + Dizel)","Kapsam 1"),
        ("1.4","Soğutucu Gaz & Yangın Söndürme","Kapsam 1"),
        ("2.1","Elektrik Tüketimi (Konum Temelli)","Kapsam 2"),
        ("3.3","Personel İşe Gidiş-Geliş","Kapsam 3"),
        ("3.5","İş Seyahatleri","Kapsam 3"),
        ("4.1","Satın Alınan Ürün & Malzemeler","Kapsam 3"),
        ("4.2","Sermaye Varlıkları (CAPEX)","Kapsam 3"),
        ("4.3","Atık Bertarafı","Kapsam 3"),
        ("4.5","Hizmet Alımları","Kapsam 3"),
        ("6.1","T&D Kayıpları","Kapsam 3"),
    ]
    add_caption(doc, "Tablo 9. Yıllara Göre Kategori Bazlı GHG Emisyonları (ton CO₂e)")
    yc = _year_cols()
    cw = [1.0, 1.8, 5.0] + [yc]*len(years) + [2.0]
    tbl = make_table(doc, len(detail_rows)+2, 3+len(years)+1, cw)
    hdrs = ["Sınıf","Kapsam","Emisyon Kaynağı"] + years + [f"Değişim ({years[0]}→{latest_year})"]
    for i, txt in enumerate(hdrs):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==2 else "center")

    for ri, (k, lbl, kapsam) in enumerate(detail_rows):
        bg = alt_bg(ri); row = tbl.rows[ri+1]
        cell_write(row.cells[0], k,      align="center", bg=bg)
        cell_write(row.cells[1], kapsam, align="center", bg=bg, sz=8, color=MID_BLUE)
        cell_write(row.cells[2], lbl,    bg=bg)
        for yi, y in enumerate(years):
            cell_write(row.cells[3+yi], fmtN(cats.get(y,{}).get(k,0)),
                       align="right", bg=bg)
        fv = cats.get(years[0],{}).get(k,0)
        lv = cats.get(latest_year,{}).get(k,0)
        chg = pct_change(lv, fv)
        col = pct_color(lv, fv)
        cell_write(row.cells[3+len(years)], chg, align="center", bg=bg, color=col)

    ft = scope_totals.get(years[0],{}).get("total",0)
    lt = scope_totals.get(latest_year,{}).get("total",0)
    gc = [("","left"),("","left"),("GENEL TOPLAM","left")]
    for y in years: gc.append((fmtN(scope_totals.get(y,{}).get("total",0)),"right"))
    gc.append((pct_change(lt,ft),"center"))
    grand_total_cells(tbl.rows[-1], gc)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# BÖLÜM 5: GHG EMİSYON BEYANI (Eski rapor formatıyla birebir — 3 ayrı tablo)
# ═══════════════════════════════════════════════════════════════════
def add_ghg_declaration(doc):
    """
    ISO 14064-1 standart GHG emisyon beyanı — eski rapordaki gibi 3 ayrı tablo:
    1. Doğrudan Emisyonlar (Kategori 1)
    2. Enerji Dolaylı Emisyonlar (Kategori 2)
    3. Diğer Dolaylı Emisyonlar (Kategori 3-4-5-6)
    + Uzaklaştırmalar + Sözleşmeye Dayalı Araçlar
    Sütunlar: Kategori | Emisyonlar | TOPLAM (CO2e) | CO2 | CH4 | N2O | HFCs | PFCs | SF6
    Yıl sırası: tüm yıllar soldan sağa (2023 → 2024 → 2025)
    """
    add_h1(doc, "BÖLÜM 5: GHG EMİSYON BEYANI")

    # Kuruluş bilgileri
    add_caption(doc, "Tablo 11. Kuruluş ve Raporlama Bilgileri")
    info_rows = [
        ("Raporlamayı Yapan Kuruluş",   "UNSPED Gümrük Müşavirliği ve Lojistik Hizmetler A.Ş."),
        ("Rapordan Sorumlu Birimler",   "UGM Test ve Analiz Laboratuvarı Departmanı"),
        ("Raporlama Periyodu",          latest_year),
        ("Kuruluş Sınırları",           "UNSPED Mahmutbey Genel Merkezi — Operasyonel Kontrol Yaklaşımı"),
        ("Raporlama Sınırları",         "Mahmutbey Genel Merkezi ve tüm operasyonel kontrolündeki faaliyetler"),
    ]
    tbl_i = make_table(doc, len(info_rows)+1, 2, [5.5, 10.3])
    hdr(tbl_i.rows[0].cells[0], "Alan", "left")
    hdr(tbl_i.rows[0].cells[1], "Bilgi", "left")
    for ri, (k, v) in enumerate(info_rows):
        bg = alt_bg(ri)
        cell_write(tbl_i.rows[ri+1].cells[0], k, bold=True, bg=bg)
        cell_write(tbl_i.rows[ri+1].cells[1], v, bg=bg)
    add_space(doc)

    # ── Yardımcı: tek yıl için standart 8 sütunlu satır ──────────
    # Sütunlar: Kategori | Emisyonlar | TOPLAM(CO2e) | CO2 | CH4 | N2O | HFCs | PFCs | SF6
    # col_widths toplam = 15.8 cm (A4 margin)
    CW = [1.0, 5.4, 1.7, 1.5, 1.3, 1.3, 1.5, 1.5, 0.6]
    HDR = ["Kategori", "Emisyonlar", "TOPLAM (CO2e)", "CO2 (CO2e)", "CH4 (CO2e)", "N2O (CO2e)", "HFCs", "PFCs", "SF6"]

    def make_decl_table(doc, title, caption_text):
        add_h2(doc, title)
        add_caption(doc, caption_text)
        tbl = make_table(doc, 1, 9, CW)
        for ci, txt in enumerate(HDR):
            hdr(tbl.rows[0].cells[ci], txt, "left" if ci==1 else "center")
        return tbl

    def add_section_row(tbl, kat, lbl, toplam, co2=None, ch4=None, n2o=None,
                        hfc=None, pfc=None, sf6=None, od="", bg=LBL_BG, is_grand=False):
        row = tbl.add_row()
        if is_grand:
            for ci, (txt, al) in enumerate([
                (kat,"center"),(lbl,"left"),(toplam,"right"),
                (co2 or "—","right"),(ch4 or "—","right"),(n2o or "—","right"),
                (hfc or "—","right"),(pfc or "—","right"),(sf6 or "—","right")
            ]):
                cell_write(row.cells[ci], txt, bold=True, color=WHITE,
                           align=al, bg=HDR_BG)
        else:
            cell_write(row.cells[0], kat,    bold=bool(bg==LBL_BG), align="center",
                       bg=bg, color=DARK_BLUE if bg==LBL_BG else None)
            cell_write(row.cells[1], lbl,    bold=bool(bg==LBL_BG), bg=bg,
                       color=DARK_BLUE if bg==LBL_BG else None)
            cell_write(row.cells[2], toplam or "—", bold=bool(bg==LBL_BG),
                       align="right", bg=bg)
            for ci, val in enumerate([co2,ch4,n2o,hfc,pfc,sf6], 3):
                cell_write(row.cells[ci], val or "—", align="right", bg=bg, sz=8)

    def add_data_row(tbl, kat, lbl, toplam, co2=None, ch4=None, n2o=None,
                     hfc=None, pfc=None, sf6=None, od="", ri=0):
        bg = alt_bg(ri)
        row = tbl.add_row()
        cell_write(row.cells[0], kat, align="center", bg=bg, sz=8,
                   color=MID_BLUE if kat else None)
        # od flag varsa etikete ekle
        lbl_text = f"{lbl}  [{od}]" if od else lbl
        cell_write(row.cells[1], lbl_text, bg=bg, sz=9,
                   italic=bool(od), color=GREY_TXT if od else None)
        cell_write(row.cells[2], toplam or "—", align="right", bg=bg)
        for ci, val in enumerate([co2,ch4,n2o,hfc,pfc,sf6], 3):
            cell_write(row.cells[ci], val or "—", align="right", bg=bg, sz=8)

    def fv(key): return fmtN(latest_cats.get(key,0))
    def fk(k):
        """Kısmi CO2/CH4/N2O ayrıştırma — yakıt yanma için yaklaşık"""
        return fmtN(latest_cats.get(k,0))

    # ══════════════════════════════════════════════════════════════
    # TABLO A: DOĞRUDAN EMİSYONLAR 2024
    # ══════════════════════════════════════════════════════════════
    k1_tot = fmtN(latest_d.get("k1",0))
    s11    = fmtN(latest_cats.get("1.1",0))
    s12    = fmtN(latest_cats.get("1.2",0))
    s14    = fmtN(latest_cats.get("1.4",0))

    tbl_a = make_decl_table(doc,
        f"Doğrudan Emisyonlar {latest_year}",
        f"Tablo 12. Doğrudan Emisyonlar {latest_year} (ton CO₂e)")

    add_section_row(tbl_a, "1",
        "Doğrudan Sera Gazı Emisyonları ve Uzaklaştırmaları (CO2e)",
        k1_tot, bg=LBL_BG)

    add_data_row(tbl_a, "1.1",
        "Sabit Yakma Kaynaklı Doğrudan Emisyonlar",
        s11, co2=s11, ri=0)
    add_data_row(tbl_a, "1.2",
        "Hareketli Yakma Kaynaklı Doğrudan Emisyonlar",
        s12, co2=s12, ri=1)
    add_data_row(tbl_a, "1.3",
        "Endüstriyel Süreçlerden Kaynaklanan Doğrudan Proses Emisyonları",
        "—", od="UY", ri=2)
    add_data_row(tbl_a, "1.4",
        "Antropojenik Sistemlerden Kaynaklı Kaçak Emisyonlar",
        s14, hfc=s14, ri=3)
    add_data_row(tbl_a, "1.5",
        "Arazi Kullanımı (LULUCF) Faaliyetleri", "—", od="UY", ri=4)
    add_data_row(tbl_a, "",
        "Biyokütle Kaynaklı Doğrudan Emisyonlar (CO2e)", "—", ri=5)

    add_section_row(tbl_a, "", "KAPSAM 1 TOPLAM", k1_tot, is_grand=True)
    add_space(doc)

    # ══════════════════════════════════════════════════════════════
    # TABLO B: ENERJİ DOLAYLI EMİSYONLAR
    # ══════════════════════════════════════════════════════════════
    k2_tot = fmtN(latest_d.get("k2",0))
    s21    = fmtN(latest_cats.get("2.1",0))

    tbl_b = make_decl_table(doc,
        f"Enerji Dolaylı Emisyonlar {latest_year}",
        f"Tablo 13. Enerji Dolaylı Emisyonlar {latest_year} (ton CO₂e)")

    add_section_row(tbl_b, "2",
        "İthal Edilen Enerjiden Kaynaklanan Dolaylı Sera Gazı Emisyonları",
        k2_tot, bg=LBL_BG)
    add_data_row(tbl_b, "2.1",
        "İthal Edilen Elektrikten Kaynaklı Dolaylı Emisyonlar (Konum Temelli)",
        s21, co2=s21, ri=0)
    add_data_row(tbl_b, "2.2",
        "İthal Edilen Diğer Enerjiden Kaynaklı Dolaylı Emisyonlar",
        "—", od="UY", ri=1)

    add_section_row(tbl_b, "", "KAPSAM 2 TOPLAM", k2_tot, is_grand=True)
    add_space(doc)

    # ══════════════════════════════════════════════════════════════
    # TABLO C: DİĞER DOLAYLI EMİSYONLAR
    # ══════════════════════════════════════════════════════════════
    k3_tot = fmtN(latest_d.get("k3",0))
    s33 = fmtN(latest_cats.get("3.3",0))
    s35 = fmtN(latest_cats.get("3.5",0))
    s41 = fmtN(latest_cats.get("4.1",0))
    s42 = fmtN(latest_cats.get("4.2",0))
    s43 = fmtN(latest_cats.get("4.3",0))
    s44 = fmtN(latest_cats.get("4.4",0))
    s45 = fmtN(latest_cats.get("4.5",0))
    s61 = fmtN(latest_cats.get("6.1",0))
    k3_transport = fmtN(latest_cats.get("3.3",0) + latest_cats.get("3.5",0))
    k4_all = fmtN(sum(latest_cats.get(k,0) for k in ["4.1","4.2","4.3","4.4","4.5"]))

    tbl_c = make_decl_table(doc,
        f"Diğer Dolaylı Emisyonlar {latest_year}",
        f"Tablo 14. Diğer Dolaylı Emisyonlar {latest_year} (ton CO₂e)")

    # Kategori 3
    add_section_row(tbl_c, "3",
        "Ulaşım Kaynaklı Dolaylı Sera Gazı Emisyonları", k3_transport, bg=LBL_BG)
    add_data_row(tbl_c, "3.1",
        "Mal (Kuruluşa Gelen) Taşımacılığı", "—", od="VY", ri=0)
    add_data_row(tbl_c, "3.2",
        "Mal (Kuruluştan Giden) Taşımacılığı", "—", od="VY", ri=1)
    add_data_row(tbl_c, "3.3",
        "Personelin İşe Gidiş-Gelişleri Kaynaklı Emisyonlar",
        s33, co2=s33, ri=2)
    add_data_row(tbl_c, "3.4",
        "Müşteriler ve Ziyaretçilerin Ulaşımı", "—", od="VY", ri=3)
    add_data_row(tbl_c, "3.5",
        "İş Seyahatleri Kaynaklı Emisyonlar", s35, ri=4)

    # Kategori 4
    add_section_row(tbl_c, "4",
        "Kuruluş Tarafından Kullanılan Ürünler Kaynaklı Dolaylı Emisyonlar",
        k4_all, bg=LBL_BG)
    add_data_row(tbl_c, "4.1",
        "Satın Alınan Hammadde/Mamul/Yarı Mamul Kaynaklı Emisyonlar",
        s41, ri=0)
    add_data_row(tbl_c, "4.2",
        "Sermaye Niteliğindeki Varlıklardan Kaynaklanan Emisyonlar",
        s42, ri=1)
    add_data_row(tbl_c, "4.3",
        "Katı ve Sıvı Atıkların Bertarafı Kaynaklı Emisyonlar",
        s43 if s43 != "—" else "—",
        od="" if s43 != "—" else "OD", ri=2)
    add_data_row(tbl_c, "4.4",
        "Kiralanan Ekipmanların Kullanımı Kaynaklı Emisyonlar",
        s44 if s44 != "—" else "—",
        od="" if s44 != "—" else "OD", ri=3)
    add_data_row(tbl_c, "4.5",
        "Danismanlik, Temizlik, Bakim, Kurye vb. Hizmet Alimlari",
        s45, ri=4)

    # Kategori 5
    add_section_row(tbl_c, "5",
        "Ürünlerin Üretim Sonrası Kullanımı Kaynaklı Dolaylı Emisyonlar",
        "—", od="UY", bg=LBL_BG)
    add_data_row(tbl_c, "5.1", "Ürünün Kullanımı Kaynaklı Emisyonlar", "—", od="UY", ri=0)
    add_data_row(tbl_c, "5.2", "Kiraya Verilen Ekipmanların Kullanımı", "—", od="UY", ri=1)
    add_data_row(tbl_c, "5.3", "Ürün Kullanım Ömrü Sonrası Emisyonlar", "—", od="UY", ri=2)
    add_data_row(tbl_c, "5.4", "Yatırımlar Kaynaklı Emisyonlar", "—", od="UY", ri=3)

    # Kategori 6
    add_section_row(tbl_c, "6",
        "Diğer Kaynaklardan Ortaya Çıkan Dolaylı Emisyonlar", s61, bg=LBL_BG)
    add_data_row(tbl_c, "6.1",
        "Elektrik İletim ve Dağıtım Kayıpları (Konum Temelli)", s61, co2=s61, ri=0)

    add_section_row(tbl_c, "", "KAPSAM 3 TOPLAM", k3_tot, is_grand=True)
    add_space(doc)

    # Not satırı
    add_note(doc,
        "OD = Onemli Degil (onemlilik esligi altinda)  |  "
        "VY = Veri Yok  |  UY = Uygulanamaz")
    add_space(doc)

    # ══════════════════════════════════════════════════════════════
    # UZAKLAŞTIRMALAR
    # ══════════════════════════════════════════════════════════════
    add_h2(doc, "Uzaklaştırmalar")
    add_caption(doc, "Tablo 15. Uzaklaştırmalar (ton CO₂e)")
    tbl_u = make_table(doc, 3, 9, CW)
    for ci, txt in enumerate(HDR): hdr(tbl_u.rows[0].cells[ci], txt, "left" if ci==1 else "center")
    for ri, lbl in enumerate(["Doğrudan Uzaklaştırmalar"]):
        add_data_row(tbl_u, ri+1, lbl, "—", ri=ri)
    # Başlık satırı için özel
    row_u = tbl_u.rows[1]
    cell_write(row_u.cells[0], "1", align="center", bg=alt_bg(0))
    cell_write(row_u.cells[1], "Doğrudan Uzaklaştırmalar (CO2e)", bg=alt_bg(0))
    for ci in range(2,9): cell_write(row_u.cells[ci], "—", align="right", bg=alt_bg(0), sz=8)
    row_u2 = tbl_u.rows[2]
    cell_write(row_u2.cells[0], "", align="center", bg=alt_bg(1))
    cell_write(row_u2.cells[1], "Biyokütle Kaynaklı Uzaklaştırmalar (CO2e)", bg=alt_bg(1))
    for ci in range(2,9): cell_write(row_u2.cells[ci], "—", align="right", bg=alt_bg(1), sz=8)
    add_space(doc)

    # ══════════════════════════════════════════════════════════════
    # SÖZLEŞMEYE DAYALI ARAÇLAR
    # ══════════════════════════════════════════════════════════════
    add_h2(doc, "Sözleşmeye Dayalı Araçlar")
    add_caption(doc, "Tablo 16. Sözleşmeye Dayalı Araçlar")
    tbl_s = make_table(doc, 4, 5, [5.5, 2.0, 2.0, 2.0, 2.0])
    for ci, txt in enumerate(["Araç","Miktar (kW)",
                               "ISO 14064-1 Ek-E Uyumu",
                               "EF (gCO2e/kW)","Emisyon (tCO2e)"]):
        hdr(tbl_s.rows[0].cells[ci], txt, "left" if ci==0 else "center")
    for ri, (lbl, uyum) in enumerate([
        ("Satın Alınan Toplam Yenilenebilir Enerji", "—"),
        ("A Tedarikçisi — Sözleşmeli Yenilenebilir Enerji", "Evet"),
        ("B Tedarikçisi — Sözleşmeli Yenilenebilir Enerji", "Hayır"),
    ]):
        bg = alt_bg(ri); row = tbl_s.rows[ri+1]
        cell_write(row.cells[0], lbl, bg=bg)
        cell_write(row.cells[1], "—", align="center", bg=bg)
        cell_write(row.cells[2], uyum, align="center", bg=bg)
        cell_write(row.cells[3], "—", align="center", bg=bg)
        cell_write(row.cells[4], "—", align="center", bg=bg)

    add_note(doc,
        "ISO 14064-1:2018 EK-E'ye uygun bir satin alma soz konusu olmadigi "
        "zaman EF hucre bos birakilacaktir.")
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SONUÇ
# ═══════════════════════════════════════════════════════════════════
def add_conclusion(doc):
    add_h1(doc, "BÖLÜM 6: SONUÇ")
    t = latest_d.get("total",0); p = prev_d.get("total",0)
    chg = pct_change(t,p) if prev_year else ""
    add_body(doc, f"{latest_year} toplam emisyonu {fmtN(t)} tCO2e olarak hesaplanmistir." +
             (f" {prev_year}'e gore {chg} degisim gostermistir." if prev_year else ""))
    add_space(doc)
    add_caption(doc, "Tablo 17. Emisyon Sonuc Tablosu")
    yc = min(2.5, 9.0/max(len(years),1))
    cw = [4.5, 1.5] + [yc]*len(years)
    tbl = make_table(doc, 5, 2+len(years), cw)
    for i, txt in enumerate(["Kapsam","Birim"]+sorted(years)):
        hdr(tbl.rows[0].cells[i], txt, "left" if i==0 else "center")
    for ri, (key, lbl, bg) in enumerate([
        ("k1","Kapsam 1 — Dogrudan Emisyonlar",      LBL_BG),
        ("k2","Kapsam 2 — Enerji Dolayli Emisyonlar", GRN_BG),
        ("k3","Kapsam 3 — Diger Dolayli Emisyonlar",  ORG_BG2),
    ]):
        row = tbl.rows[ri+1]
        cell_write(row.cells[0], lbl, bold=True, bg=bg, color=DARK_BLUE)
        cell_write(row.cells[1], "tCO2e", align="center", bg=bg, sz=8)
        for yi, y in enumerate(sorted(years)):
            cell_write(row.cells[2+yi], fmtN(scope_totals.get(y,{}).get(key,0)),
                       align="right", bg=bg)
    gc = [("GENEL TOPLAM","left"),("tCO2e","center")]
    for y in sorted(years): gc.append((fmtN(scope_totals.get(y,{}).get("total",0)),"right"))
    grand_total_cells(tbl.rows[4], gc)
    add_space(doc)


# ═══════════════════════════════════════════════════════════════════
# ÖNERILER
# ═══════════════════════════════════════════════════════════════════
def add_recommendations(doc):
    add_h1(doc, "BOLUM 7: ONERILER")
    add_body(doc, "Emisyon azaltimi icin asagidaki stratejik oncelikler belirlenmistir:")
    add_space(doc)
    add_h2(doc, "7.1. Emisyon Azaltimi")
    for konu, aciklama in [
        ("Hizmet Alimlari (4.5)", "Dusuk/sifir emisyonlu tasimacilarin tercih edilmesi."),
        ("Satin Alinan Urunler (4.1)", "Geri donusturulmus icerikli urunlere yonelim."),
        ("Mobil Yanma (1.2)", "Arac filosunda hibrit/elektrikli modellerin tercihi."),
        ("Elektrik (2.1)", "LED donusumu, otomasyon ve cati GES projeleri."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{konu}: ")
        r1.font.bold = True; r1.font.color.rgb = MID_BLUE; r1.font.size = Pt(10)
        r2 = p.add_run(aciklama); r2.font.size = Pt(10)
    add_h2(doc, "7.2. SBTi Uyumlu Hedefleme")
    for m in [
        "Kapsam 1-2: Yillik %4,2 azaltim (10 yilda %42)",
        "Kapsam 3: Yillik %2,5 azaltim (10 yilda %25)",
        "Yenilenebilir elektrik kullanimi: Her yil %10 artis",
        "Arac filosu elektrifikasyonu icin yillik hedefler",
    ]:
        add_bullet(doc, m)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# METODOLOJI & KAYNAKLAR
# ═══════════════════════════════════════════════════════════════════
def add_methodology(doc):
    add_h1(doc, "BOLUM 8: METODOLOJI VE REFERANS TABLOLAR")
    add_body(doc, "Bu bolumde hesaplamalarda kullanilan tum emisyon faktorleri, "
             "net kalorifik degerler ve GWP degerleri bir arada sunulmaktadir.")

    add_h2(doc, "8.1. Kullanilan Kaynaklar")
    tbl = make_table(doc, 9, 3, [0.8, 4.5, 10.5])
    for i, txt in enumerate(["#","Konu","Kaynak"]):
        hdr(tbl.rows[0].cells[i], txt, "left" if i>0 else "center")
    for ri, (n,k,v) in enumerate([
        ("1","Uluslararasi Standart","TS EN ISO 14064-1:2019"),
        ("2","Metodoloji","GHG Protocol Corporate Standard (WRI/WBCSD)"),
        ("3","Yakit EF & NCV","IPCC 2006 Guidelines Vol.2, Table 2.2"),
        ("4","GWP Degerleri","IPCC AR6 2021, Table 7.SM.7"),
        ("5","Elektrik Grid EF","TC ETKB EVCEN Emisyon Faktorleri"),
        ("6","Tasimacilk EF","UK DEFRA GHG Conversion Factors"),
        ("7","Tedarik Zinciri EF","US EPA USEEIO v1.2 (NAICS-6, 2021 USD)"),
        ("8","Sogutma Gazi GWP","IPCC AR6, Tablo 7.SM.7"),
    ]):
        bg = alt_bg(ri); row = tbl.rows[ri+1]
        cell_write(row.cells[0], n, align="center", bg=bg)
        cell_write(row.cells[1], k, bold=True, bg=bg)
        cell_write(row.cells[2], v, bg=bg)

    add_h2(doc, "8.2. Emisyon Faktorleri Tablosu")
    add_caption(doc, "Tablo EF-1. Kapsam 1 & 2 Emisyon Faktorleri")
    ef_data = [
        ("Dogalgaz (Sabit Yanma)",    "Kg/TJ","56.100","5","0,1","—","IPCC 2006 GL"),
        ("Mazot (Sabit Yanma)",        "Kg/TJ","74.100","10","0,6","—","IPCC 2006 GL"),
        ("Mazot (Mobil Yanma)",        "Kg/TJ","74.100","3,9","3,9","—","IPCC 2006 GL"),
        ("Benzin (Sabit Yanma)",       "Kg/TJ","69.300","10","0,6","—","IPCC 2006 GL"),
        ("Benzin (Mobil Yanma)",       "Kg/TJ","69.300","25","8","—","IPCC 2006 GL"),
        ("Klima Sog. Gazi (Sizinti)",  "%","—","—","—","10%","IPCC AR6"),
        ("Klima Sog. Gazi (Sarj)",     "%","—","—","—","100%","IPCC AR6"),
        ("Sebil/Buz Dolabi (Sizinti)", "%","—","—","—","0,5%","IPCC AR6"),
        ("Yangin Tupu (Desarj/Sarj)",  "%","—","—","—","100%","IPCC AR6"),
        ("Yangin Sondurme (Sizinti)",  "%","—","—","—","2%","IPCC AR6"),
        ("Elektrik (Konum Temelli)", "tCO2/MWh","0,434-0,442","—","—","—","TC ETKB EVCEN"),
        ("T&D Kayiplari",            "tCO2/MWh","Pazar-Konum","—","—","—","TC ETKB EVCEN"),
    ]
    cw_ef = [4.5, 2.0, 2.0, 1.5, 1.5, 1.8, 3.5]
    tbl_ef = make_table(doc, len(ef_data)+1, 7, cw_ef)
    for i, txt in enumerate(["Emisyon Kaynagi","Birim","EF CO2","EF CH4","EF N2O","Kutle%","Kaynak"]):
        hdr(tbl_ef.rows[0].cells[i], txt, "left" if i==0 else "center")
    for ri, row_d in enumerate(ef_data):
        bg = alt_bg(ri); row = tbl_ef.rows[ri+1]
        for ci, val in enumerate(row_d):
            cell_write(row.cells[ci], val, align="left" if ci==0 else "center", bg=bg)

    add_h2(doc, "8.3. Net Kalorifik Deger (NKD) Tablosu")
    add_caption(doc, "Tablo EF-2. Yakit NKD ve Yogunluk Degerleri (IPCC 2006 GL)")
    ncv_data = [
        ("Dogalgaz","m3","48,0","TJ/Gg","0,000717","ton/m3"),
        ("Dizel/Mazot","L","43,0","TJ/Gg","0,000832","ton/L"),
        ("Benzin","L","44,3","TJ/Gg","0,000745","ton/L"),
        ("LPG","L","47,3","TJ/Gg","0,000540","ton/L"),
        ("Fuel Oil","L","40,4","TJ/Gg","0,000890","ton/L"),
        ("Komur","ton","11,9","TJ/Gg","1,0","ton/ton"),
    ]
    tbl_n = make_table(doc, len(ncv_data)+1, 6, [3.5, 1.5, 2.0, 2.5, 2.5, 3.8])
    for i, txt in enumerate(["Yakit","Birim","NKD","NKD Birimi","Yogunluk","Yogunluk Birimi"]):
        hdr(tbl_n.rows[0].cells[i], txt, "left" if i==0 else "center")
    for ri, row_d in enumerate(ncv_data):
        bg = alt_bg(ri); row = tbl_n.rows[ri+1]
        for ci, val in enumerate(row_d):
            cell_write(row.cells[ci], val, align="left" if ci==0 else "center", bg=bg)

    add_h2(doc, "8.4. GWP-100 Degerleri (IPCC AR6 2021)")
    add_caption(doc, "Tablo EF-3. Sogutma Gazi GWP-100 Degerleri")
    gwp_data = [
        ("R410A","HFC-410A","2.088","IPCC AR6"),
        ("R407c","HFC-407C","1.774","IPCC AR6"),
        ("R134A","HFC-134a","1.526","IPCC AR6"),
        ("R32","HFC-32","771","IPCC AR6"),
        ("HFC236fa","HFC-236fa","8.060","IPCC AR6"),
        ("FM200/HFC-227ea","HFC-227ea","3.220","IPCC AR6"),
        ("R600A","Izobutan","4","IPCC AR6"),
        ("FK-5-1-12","Novec 1230","< 1","IPCC AR6"),
    ]
    tbl_g = make_table(doc, len(gwp_data)+1, 4, [2.5, 6.0, 2.5, 3.8])
    for i, txt in enumerate(["Formul","Gaz Adi","GWP-100","Kaynak"]):
        hdr(tbl_g.rows[0].cells[i], txt, "left" if i<=1 else "center")
    for ri, (f,n,g,s) in enumerate(gwp_data):
        bg = alt_bg(ri); row = tbl_g.rows[ri+1]
        cell_write(row.cells[0], f, bold=True, bg=bg)
        cell_write(row.cells[1], n, bg=bg)
        cell_write(row.cells[2], g, align="right", bg=bg)
        cell_write(row.cells[3], s, align="center", bg=bg)

    add_h1(doc, "BOLUM 9: KISALTMALAR VE TANIMLAR")
    kisalt = [
        ("CO2e","Karbon Dioksit Esdegeri"),("GHG","Sera Gazlari"),
        ("GWP","Kuresel Isinma Potansiyeli"),("IPCC","Hukumetlerarasi Iklim Paneli"),
        ("DEFRA","BK Cevre Bakanligi"),("EPA","ABD Cevre Koruma Ajansi"),
        ("USEEIO","ABD Cevre Ekonomik Girdi-Cikti Modeli"),
        ("EVCEN","TC ETKB Enerji Verimliligi ve Cevre"),
        ("NCV/NKD","Net Kalorifik Deger"),("HFC","Hidroflorokarbon"),
        ("T&D","Iletim ve Dagitim"),("SBTi","Bilim Bazli Hedefler Girisimi"),
        ("CAPEX","Sermaye Nitelikli Harcamalar"),("tCO2e","Ton CO2 Esdegeri"),
        ("EF","Emisyon Faktoru"),("FV","Faaliyet Verisi"),
    ]
    tbl2 = make_table(doc, len(kisalt)+1, 2, [3.5, 13.3])
    hdr(tbl2.rows[0].cells[0], "Kisaltma", "left")
    hdr(tbl2.rows[0].cells[1], "Aciklama", "left")
    for ri, (k,v) in enumerate(kisalt):
        bg = alt_bg(ri)
        cell_write(tbl2.rows[ri+1].cells[0], k, bold=True, bg=bg)
        cell_write(tbl2.rows[ri+1].cells[1], v, bg=bg)


# ═══════════════════════════════════════════════════════════════════
# HEADER & FOOTER
# ═══════════════════════════════════════════════════════════════════
def add_header_footer(doc):
    sec = doc.sections[0]
    header = sec.header
    para   = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"2E75B6")
    pBdr.append(bot); pPr.append(pBdr)
    r1 = para.add_run("UNSPED")
    r1.font.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = DARK_BLUE
    r2 = para.add_run(f"  |  Kurumsal Karbon Ayak Izi Envanter Raporu  |  {latest_year}")
    r2.font.size = Pt(9); r2.font.color.rgb = GREY_TXT
    footer = sec.footer
    para2  = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para2.clear()
    pPr2 = para2._p.get_or_add_pPr(); pBdr2 = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),"single"); top.set(qn("w:sz"),"6")
    top.set(qn("w:space"),"1"); top.set(qn("w:color"),"2E75B6")
    pBdr2.append(top); pPr2.append(pBdr2)
    rf = para2.add_run(
        f"TS EN ISO 14064-1:2019  |  GHG Protocol  |  UNSPED {latest_year}  |  Sayfa ")
    rf.font.size = Pt(8); rf.font.color.rgb = GREY_TXT
    for fld_type in ["begin","end"]:
        run_el = OxmlElement("w:r")
        rpr    = OxmlElement("w:rPr")
        sz_el  = OxmlElement("w:sz"); sz_el.set(qn("w:val"),"16")
        rpr.append(sz_el); run_el.append(rpr)
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), fld_type)
        if fld_type == "begin":
            run_el.append(fc)
            instr = OxmlElement("w:r")
            it    = OxmlElement("w:instrText"); it.text = "PAGE"
            instr.append(it)
            para2._p.append(run_el); para2._p.append(instr)
        else:
            run_el.append(fc); para2._p.append(run_el)


# ═══════════════════════════════════════════════════════════════════
# ANA AKIŞ
# ═══════════════════════════════════════════════════════════════════
add_cover(doc)
if "intro"            in selected: add_intro(doc)
if "technical"        in selected: add_technical(doc)
if "boundaries"       in selected: add_boundaries(doc)
if "scope1"           in selected: add_scope1(doc)
if "scope2"           in selected: add_scope2(doc)
if "scope3"           in selected: add_scope3(doc)
if "activity"         in selected: add_activity_data(doc)
if "charts"           in selected: add_charts(doc)
if "summary"          in selected: add_summary(doc)
if "impact"           in selected: add_impact(doc)
if "percap"           in selected: add_percap(doc)
if "yearly"           in selected: add_yearly(doc)
if "ghg_declaration"  in selected: add_ghg_declaration(doc)
if "conclusion"       in selected: add_conclusion(doc)
if "recommendations"  in selected: add_recommendations(doc)
if "methodology"      in selected: add_methodology(doc)
add_header_footer(doc)
doc.save(out_path)
print(f"OK:{out_path}")